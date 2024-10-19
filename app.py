import streamlit as st
from PIL import Image
import numpy as np
import tensorflow as tf
import base64
import google.generativeai as genai
import io
import docx
import os
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

# Load your pre-trained model (replace 'your_model.h5' with the actual path to your model)
model = tf.keras.models.load_model('model.h5')

# Define a function to preprocess the image
def preprocess_image(image, target_size=(224, 224)):
    image = image.resize(target_size)
    image = np.array(image)
    if image.shape[2] == 4:  # if RGBA, convert to RGB
        image = image[..., :3]
    image = image / 255.0  # Normalize the image to [0, 1]
    image = np.expand_dims(image, axis=0)  # Add batch dimension
    return image

# Define a function to make predictions
def predict(image, model):
    processed_image = preprocess_image(image)
    result = model.predict(processed_image)
    result = result.ravel() 
    classes = ["Fusarium Wilt", "Leaf Curl Disease", "Healthy Leaf", "Healthy Plant"]
    max = result[0];    
    index = 0; 
    #Loop through the array    
    for i in range(0, len(result)):    
      #Compare elements of array with max    
      if(result[i] > max):    
          max = result[i];    
          index = i
    #print("Largest element present in given array: " + str(max) +" And it belongs to " +str(classes[index]) +" class."); 
    pred = str(classes[index])
    return pred

def img_to_base64(image_path):
    """Convert image to base64."""
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except Exception as e:
        return None
    
# Configure Google Generative AI
genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))

# Set up the model
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 0,
    "max_output_tokens": 8192,
}

safety_settings = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
]

system_prompts = {
    "image": """
    You are a domain expert in agriculture image analysis. You are tasked with 
    examining cotton disease images for a renowned indian agriculture farming. Your expertise will help in identifying 
    or discovering any anomalies, diseases, conditions or any planty related issues that might be present in the image.
    
    Your key responsibilities:
    1. Detailed Analysis: Scrutinize and thoroughly examine each image, 
       focusing on finding any abnormalities.
    2. Analysis Report: Document all the findings and 
       clearly articulate them in a structured format.
    3. Recommendations: Based on the analysis, suggest remedies, 
       tests or treatments as applicable.
    4. Treatments: If applicable, lay out detailed treatments 
       which can help in faster recovery.
    
    Important Notes to remember:
    1. Scope of response: Only respond if the image pertains to plant related issues.
    2. Clarity of image: In case the image is unclear, 
       note that certain aspects are 'Unable to be correctly determined based on the uploaded image'
    3. Disclaimer: Accompany your analysis with the disclaimer: 
       "This is an AI BOT made by Indian Agriculture Farming. Consult with a agriculture expert before making any decisions."
    
    Please provide the final response in headings and sub-headings in bullet format: 
    Detailed Analysis, Analysis Report, Recommendations and Treatments.

    when you found \n or \n\n in responses extend the output in next line.

    Note: If images are not related to agriculture topics, as you're a Agriculture AI Chatbot, 
    please inform the user that you can only analyze agriculture-related images.
    """,

    "text": """
    You are an AI agriculture assistant designed to provide concise, accurate information on various cotton disease plants topics. Your role is to offer brief, helpful responses to cotton-disease-related queries, including suggestions appropriate.

    When responding to queries, structure your answer concisely with these key areas as applicable:

    1. Brief Overview: Concise information about the agriculture topic or condition.
    2. Key Symptoms: List 3-5 main symptoms if relevant.
    3. Agriculture Suggestions: 
       - Suggest 2-3 commonly used agricultures for the condition, including generic names.
       - Briefly explain the purpose of each cotton disease related precausion.
    4. Quick Suggestions: Offer 2-3 practical agricultural tips or lifestyle modifications.

    Important guidelines:
    1. Keep responses short and to the point, typically within 200-250 words.
    2. Use simple, clear language accessible to a general audience.
    

    Provide your response in a simple format with appropriate headings.

    Remember, you're an AI assistant providing general information, not a replacement for professional agriculture advice or diagnosis. Always encourage users to consult with a healthcare provider for specific medical advice and treatment.
    """
}

# Update the model initialization to use the new prompts
model_image = genai.GenerativeModel(model_name="gemini-1.5-flash",
                                    generation_config=generation_config,
                                    safety_settings=safety_settings)

model_text = genai.GenerativeModel(model_name="gemini-1.5-flash",
                                   generation_config=generation_config,
                                   safety_settings=safety_settings)

def is_xray(image):
    # This is a simple check. You might want to implement a more sophisticated method
    # to detect X-ray images, possibly using machine learning.
    return image.mode == 'L'  # Check if the image is grayscale

def generate_doc(content, is_xray_image=False):
    doc = docx.Document()
    
    # Split the content into sections
    sections = content.split('**')
    
    for i, section in enumerate(sections):
        if i % 2 == 0:  # Even indexes are normal text or empty
            continue
        else:  # Odd indexes are headers or content
            # Add the header in bold
            p = doc.add_paragraph()
            p.add_run(section.strip()).bold = True
            
            # If there's content following this header, add it as bullet points
            if i + 1 < len(sections):
                content = sections[i + 1]
                bullet_points = content.split('*')
                for point in bullet_points:
                    if point.strip():
                        doc.add_paragraph(point.strip(), style='List Bullet')
    
    # Add the disclaimer as a separate paragraph
    #disclaimer = "This is an AI BOT made by MEDI360. Consult with a Doctor before making any decisions."
    #doc.add_paragraph(disclaimer)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
    
def home():
# Streamlit app
    st.title("AI Cotton Disease Classification 🍃")

    # Upload image
    uploaded_file = st.file_uploader("Choose an image...", type=["jpg", "jpeg", "png"])

    if uploaded_file is not None:
        # Display uploaded image
        image = Image.open(uploaded_file)
        st.image(image, caption='Uploaded Image', use_column_width=True)
    
        # Make a prediction
        st.write("Classifying...")
        prediction = predict(image, model)
    
        # Display the prediction
        st.write(f"Prediction: {prediction}")
        #st.write(f"Confidence: {np.max(prediction):.2f}")    

def services():
    
    st.title("AI Cotton Disease Classification Assistant 🍃")
    st.subheader("An app to help with chest diease analysis using images and text")

    os.environ["GOOGLE_API_KEY"] == st.secrets["GOOGLE_API_KEY"]

    tab1, tab2 = st.tabs(["Image Analysis", "Text Query"])

    with tab1:
        st.header("Image Analysis")
        uploaded_file = st.file_uploader("Upload the image for Analysis:", type=["jpg", "jpeg", "png"])
        if uploaded_file is not None:
            image = Image.open(uploaded_file)
            st.image(image, caption="Uploaded Image", use_column_width=True)
            if st.button("Analyze"):
                with st.spinner("Analyzing..."):
                    image_data = uploaded_file.getvalue()
                    image_parts = [{"mime_type": "image/jpeg", "data": image_data}]
                    prompt_parts = [image_parts[0], system_prompts["image"]]
                    ai_response = model_image.generate_content(prompt_parts)
                    if ai_response:
                        st.markdown(ai_response.text)

                    
                        is_xray_image = is_xray(image)
                    
                        # Generate and offer download of .doc file
                        doc_io = generate_doc(ai_response.text)
                        st.download_button(
                            label="Download analysis as .doc",
                            data=doc_io,
                            file_name="cotton_image_analysis.doc",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("An error occurred during analysis.")

    with tab2:
        st.header("Text Query")
        query = st.text_area("Enter your medical query:")
        if st.button("Submit Query"):
            with st.spinner("Processing..."):
                prompt_parts = [
                    system_prompts["text"],
                    f"User query: {query}"
                ]
                ai_response = model_text.generate_content(prompt_parts)
                if ai_response:
                    st.markdown(ai_response.text)
                
                    # Generate and offer download of .doc file
                    doc_io = generate_doc(ai_response.text)
                    st.download_button(
                        label="Download response as .doc",
                        data=doc_io,
                        file_name="cotton_query_response.doc",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("An error occurred while processing your query.")
   

# Set the background image using CSS
page_bg_img = '''
<style>
.stApp {
    background-image: url("https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcSpcEqKQUkxVRqdRQROIVjXWn_ZPN6WCXojNw&s");
    background-size: cover;
    background-position: center;
    background-repeat: no-repeat;
}
</style>
'''

# Inject the CSS into the Streamlit app
st.markdown(page_bg_img, unsafe_allow_html=True)

# Create a sidebar for navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio("Go to", ("Home", "Services"))

st.sidebar.title("About")
st.sidebar.info("This is a AI Cotton Disease Classification Assistant app that uses AI to analyze cotton disease images and answer cotton related disease queries.")

# Display the selected page
if page == "Home":
    home()
        
elif page == "Services":
    services()
